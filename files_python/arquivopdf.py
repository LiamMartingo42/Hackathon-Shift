import docx
from docx import Document
from docx.shared import Inches, Pt
from docx2pdf import convert

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
docs = Document()
titulo = docs.add_heading('', 0)
run = titulo.add_run('Consulta de Agendamento')
font = run.font
font.name = 'Poppins'
font.size = Pt(25)
font.bold = True
titulo2 = docs.add_heading('', 1)
run1 = titulo.add_run('\nTelefone do Doutor: 17 994231499 | Médico : Dr.Daniel Bianchi')
run1.font.name = 'Poppins'
run1.font.size = Pt(16)
corpo = docs.add_paragraph().add_run('''Enderenço: Av.Philadelpha Manoel Gouveia Neto,29\nHorário: 09:00h\nHospital: Hospital Dia\n Nome: Lucas Augusto Lourenço\n\nProntuário: 10001 \n\nSetor: 1 \n\nPeso:80 Kg \tAltura: 1.8
\nDepto: 2°   \tLeito: 3\n\nDiagnóstico: \n\nColeta: Data Marcada:20/11/2022\n\n
Hemograma completo\nVHS \nLeucograma \nEritograma \nPlaquetas Hemoglobina (Hb) \nHematócrito (Ht) \nReticulócitos Prova do laço T. de sangramento (TS) \nT. de coagulação (TC) \nT. de protrombina (TP) + INR T. de tromboplastina (TTPa) \nT. de trombina (TT) \nFibrinogênio Antitrombina III \nAgregação plaquetária T. de coag. ativado (TCA) \nDímero D (DDI) \nFalcização
''')
font = corpo.font
font.name = 'Poppins'
font.size = Pt(12)

docs.save('C:/Users/liamm/PycharmProjects/pythonProject/Exame1.docx')
print(getText('C:/Users/liamm/PycharmProjects/pythonProject/Exame1.docx'))
convert("C:/Users/liamm/PycharmProjects/pythonProject/Exame1.docx", "C:/Users/liamm/PycharmProjects/pythonProject/Exame1.pdf")
convert("C:/Users/liamm/PycharmProjects/pythonProject/")